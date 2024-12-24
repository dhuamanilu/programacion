from docx import Document
from docx.shared import Inches

# Crear el documento
doc = Document()

doc.add_heading('Video Brief: How to Stay Organized at Home', level=1)

# Crear la tabla
table = doc.add_table(rows=1, cols=2)

table.style = 'Table Grid'  # Estilo de tabla

# Encabezados
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Voice'
hdr_cells[1].text = 'Visuals'

# Contenido de la tabla
data = [
    ("Have you ever looked around your house and felt overwhelmed by the mess? Maybe you don’t know where to start cleaning, and things keep getting lost?", 
     "Close-up of a messy living room with clothes, books, and dishes everywhere."),
    ("Staying organized at home can be hard, but it doesn’t have to be. Here’s how you can keep your space tidy and stress-free.", 
     "Person looking confused and frustrated, holding a pile of random objects."),
    ("The main problem is that we don’t have a system for organizing our things, so items end up in the wrong places.", 
     "A room with items scattered everywhere – shoes on the table, books on the couch, and toys on the floor, showing a lack of organization."),
    ("Here are three simple steps to help you stay organized:", 
     "Text on screen: 'Follow These Steps.'"),
    ("**Step 1**: Start small. Focus on one room or one area at a time. For example, clean your desk before tackling the whole house.", 
     "Person cleaning a desk, removing papers, and putting items into drawers."),
    ("**Step 2**: Have a place for everything. Use baskets, shelves, and labels to know where things go. This makes it easier to clean up later.", 
     "Close-up of labeled baskets and neatly organized shelves."),
    ("**Step 3**: Make it a habit. Spend 10 minutes each day organizing. Small daily actions prevent a big mess later.", 
     "Person setting a timer for 10 minutes and tidying up the room."),
    ("Staying organized at home is not about being perfect—it’s about having simple habits that make life easier.", 
     "Clean and tidy room with someone smiling and feeling relaxed."),
    ("Try these tips, and you’ll see how much more relaxed and happy you feel at home!", 
     "Final shot: Text on screen, 'A clean space = a clear mind.'")
]

# Añadir filas a la tabla
for voice, visuals in data:
    row_cells = table.add_row().cells
    row_cells[0].text = voice
    row_cells[1].text = visuals

# Guardar el documento
doc.save('video_brief_organized_home.docx')

print("Document generated successfully: video_brief_organized_home.docx")
