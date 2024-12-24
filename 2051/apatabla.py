from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_legal_framework_table():
    # Crear el documento
    doc = Document()

    # Título de la tabla
    title = doc.add_paragraph("Tabla 1\n", style='Heading 1')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.add_run("Comparativa de Marcos Legales sobre IA y Ciberseguridad").bold = True

    # Crear la tabla
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    # Encabezados
    headers = ["Región/País", "Marco Legal sobre IA", "Marco Legal sobre Ciberseguridad", "Aspectos Relevantes"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Contenido de la tabla
    data = [
        ["Unión Europea", "Reglamento de Inteligencia Artificial (RIA)", "Directiva NIS2", "Clasificación por riesgo, transparencia, resiliencia en sectores críticos"],
        ["Estados Unidos", "Orden Ejecutiva 2023", "Ley de Protección de Infraestructuras Críticas", "Enfoque en ética y protección de infraestructuras"],
        ["Perú", "Estrategia Nacional de IA (ENIA)", "Ley N.º 30999 y ENSD", "Promoción de soluciones éticas, protección de infraestructuras críticas"],
        ["China", "Ley General de IA", "Ley de Ciberseguridad", "Regulación de algoritmos, control de infraestructuras críticas"],
        ["África", "Estrategia Continental para IA", "Normas emergentes continentales", "Inclusión digital y desarrollo sostenible"]
    ]

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Nota APA debajo de la tabla
    note = doc.add_paragraph()
    note.add_run("Nota: La tabla presenta un resumen de los marcos legales relevantes en diferentes regiones y países, enfatizando la intersección entre inteligencia artificial y ciberseguridad.").italic = True

    # Guardar el archivo
    file_path = 'Comparativa_Marcos_Legales_Ampliada.docx'
    doc.save(file_path)

    return file_path

# Ejecutar la función
create_legal_framework_table()
