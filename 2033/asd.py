from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Crear documento
doc = Document()
doc.add_heading('Stakeholders del Proyecto', level=1)

# Datos de la tabla
data = [
    ["Nombre", "Puesto", "Organización / Empresa", "Ubicación", "Rol en el proyecto",
     "Información de contacto", "Requisitos principales", "Expectativas principales",
     "Grado de influencia", "Grado de interés", "Fase de mayor interés",
     "Interno / Externo", "Partidario / Neutral / Reticente", "Descripción detallada"],
    ["Yasiel Perez Vera", "Docente", "Universidad Nacional de San Agustín", "Arequipa, Perú",
     "Asesor del proyecto", "N/A", "Supervisión técnica y metodológica",
     "Asegurar el éxito académico y técnico del proyecto", "Alto", "Medio",
     "Inicio y evaluación final", "Interno", "Partidario",
     "Brinda orientación técnica y metodológica al equipo, supervisando la correcta ejecución de las actividades."],
    ["Jesús Martín Silva Fernández", "Docente de teoría", "Universidad Nacional de San Agustín",
     "Arequipa, Perú", "Apoyo teórico y conceptual", "N/A", "Proveer fundamentos teóricos",
     "Fortalecer los conocimientos técnicos del equipo", "Medio", "Bajo",
     "Planificación y diseño", "Interno", "Neutral",
     "Apoya con conceptos teóricos y fundamentos relacionados al curso, fortaleciendo el conocimiento del equipo."],
    ["Azurin Zuñiga, Eberth Wilfredo", "Jefe de Desarrollo Técnico", "N/A", "Arequipa, Perú",
     "Desarrollo Front-End y diseño de interfaces", "N/A", "Implementación de interfaces y funcionalidades visuales",
     "Interfaces funcionales y visualmente atractivas", "Muy Alto", "Alto",
     "Implementación", "Interno", "Partidario",
     "Responsable de diseñar y construir la interfaz de usuario (UI), asegurando que sea intuitiva y funcional."],
    ["Canal Mendoza, Fernando Rubén", "Gestor de Requisitos y Calidad", "N/A", "Arequipa, Perú",
     "Definición de requerimientos y pruebas de calidad", "N/A", "Requerimientos claros y pruebas exhaustivas",
     "Sistema funcional y de alta calidad", "Muy Alto", "Alto",
     "Revisión de requerimientos", "Interno", "Partidario",
     "Define los requerimientos del sistema y asegura que estos se cumplan durante el desarrollo."],
    ["Huamani Luque, Diego Alonso", "Líder del Proyecto", "N/A", "Arequipa, Perú",
     "Coordinación y diseño de base de datos", "N/A", "Gestión eficiente del equipo y bases de datos sólidas",
     "Proyecto bien gestionado y entregado a tiempo", "Muy Alto", "Muy Alto",
     "Todo el proyecto", "Interno", "Partidario",
     "Supervisa la gestión del proyecto, toma decisiones estratégicas y coordina al equipo."],
    ["Galvez Quilla, Henry Isaias", "Coordinador de Sistemas", "N/A", "Arequipa, Perú",
     "Integración de sistemas y geolocalización", "N/A", "Precisión en los servicios de geolocalización",
     "Funcionalidad precisa y bien integrada", "Muy Alto", "Alto",
     "Implementación", "Interno", "Partidario",
     "Integra los servicios de geolocalización y funcionalidad de mapas en la aplicación."],
    ["Garcia Puma, Ayrton Robins", "Coordinador de Logística", "N/A", "Arequipa, Perú",
     "Gestión de recursos y monitoreo", "N/A", "Disponibilidad de recursos y cumplimiento de cronogramas",
     "Flujo de trabajo eficiente", "Alto", "Medio",
     "Planificación", "Interno", "Partidario",
     "Planifica los recursos necesarios y monitorea el cumplimiento del cronograma."],
    ["Usuarios finales", "N/A", "Ciudadanos y turistas", "Arequipa, Perú",
     "Beneficiarios principales", "N/A", "Sistema usable y útil",
     "Experiencia de usuario óptima", "Muy Alto", "Muy Alto",
     "Pruebas y lanzamiento", "Externo", "Neutral",
     "Determinan el éxito del proyecto a través de su experiencia de usuario."],
    ["Google (API Google Maps)", "Proveedor de servicios", "Google", "Internacional",
     "Geolocalización y mapas", "N/A", "Disponibilidad y estabilidad del servicio",
     "Integración estable y funcional", "Alto", "Alto",
     "Implementación", "Externo", "Neutral",
     "Suministra la funcionalidad de mapas y geolocalización, esenciales para el sistema."],
    ["Universidad Nacional de San Agustín", "Institución educativa", "Universidad Nacional de San Agustín", "Arequipa, Perú",
     "Apoyo institucional y logístico", "N/A", "Proveer recursos, infraestructura y soporte",
     "Éxito del proyecto para cumplir objetivos académicos", "Muy Alto", "Muy Alto",
     "Todo el proyecto", "Interno", "Partidario",
     "Proporciona infraestructura, apoyo administrativo y respaldo académico."],
    ["Proveedores de APIs", "Proveedor de datos", "Variados", "Internacional",
     "Proveen datos y servicios externos", "N/A", "Calidad y disponibilidad de los datos",
     "Servicios estables y confiables", "Medio", "Medio",
     "Implementación", "Externo", "Neutral",
     "Aportan datos y servicios adicionales necesarios para la implementación."],
    ["Comunidad académica de la UNSA", "N/A", "N/A", "Arequipa, Perú",
     "Stakeholder indirecto.", "N/A", "Representada por estudiantes, docentes y personal administrativo",
     "Evaluar impacto y aplicabilidad dentro del contexto educativo", "Medio", "Medio",
     "Evaluación final", "Interno", "Neutral",
     "Representa a estudiantes, docentes y personal administrativo como evaluadores indirectos del impacto del proyecto."],
    ["Autoridades municipales de Arequipa", "N/A", "N/A", "Arequipa, Perú",
     "Potenciales socios para la implementación.", "N/A", "Fomentar el turismo y la cultura",
     "Apoyar la difusión e implementación a nivel local", "Bajo", "Medio",
     "Post-implementación", "Externo", "Neutral",
     "Pueden aprovechar el sistema para fomentar el turismo y la cultura en Arequipa."]
]

# Agregar tabla al documento
table = doc.add_table(rows=1, cols=len(data[0]))
table.style = 'Table Grid'

# Agregar encabezados
header_cells = table.rows[0].cells
for i, header in enumerate(data[0]):
    header_cells[i].text = header
    header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_cells[i].paragraphs[0].runs[0].font.bold = True
    header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

# Agregar contenido
for row_data in data[1:]:
    row_cells = table.add_row().cells
    for i, cell_data in enumerate(row_data):
        row_cells[i].text = cell_data
        row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

# Guardar documento
doc.save('Stakeholders_Proyecto_Con_Descripcion.docx')

print("Documento generado: Stakeholders_Proyecto_Con_Descripcion.docx")
