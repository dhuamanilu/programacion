from docx import Document
from docx.shared import Pt

# Datos de la tabla
data = [
    ["Nombre", "Cantidad", "Componente"],
    ["U1", 1, "Arduino Uno R3"],
    ["U2", 1, "Sensor de temperatura (TMP36)"],
    ["PIEZO1", 1, "Zumbador Piezoeléctrico"],
    ["Digit2", 1, "Display de 7 segmentos (cátodo común)"],
    ["D2", 1, "LED Verde"],
    ["D3", 1, "LED Rojo"],
    ["D4", 1, "LED Amarillo"],
    ["R1, R2, R3", 3, "Resistencia de 220 Ω"],
    ["R4, R5", 2, "Resistencia de 330 Ω"],
    ["PIR1", 1, "Sensor PIR"]
]

# Crear documento
doc = Document()
doc.add_heading("Lista de Materiales", level=1)

# Crear la tabla
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"

# Agregar encabezados
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Nombre"
hdr_cells[1].text = "Cantidad"
hdr_cells[2].text = "Componente"

# Llenar la tabla con datos
for row_data in data[1:]:
    row = table.add_row().cells
    row[0].text = str(row_data[0])
    row[1].text = str(row_data[1])
    row[2].text = row_data[2]

# Ajustar el tamaño de fuente
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(12)

# Guardar el documento
doc.save("lista_materiales.docx")
print("Documento generado: lista_materiales.docx")
