from docx import Document

# Data for risks, sorted by exposure in descending order
risk_data = [
    ["Problemas de Seguridad en la Transmisión de Datos", 0.4, 0.9, 0.36],
    ["Dificultades en la Integración con el Servicio de Geolocalización debido a Cambios en la API", 0.4, 0.9, 0.36],
    ["Pérdida de Datos en la Base de Datos Local durante Pruebas en Campo", 0.4, 0.9, 0.36],
    ["Retrasos en la Implementación del Módulo de Notificaciones de Proximidad", 0.6, 0.7, 0.42],
    ["Dependencia Excesiva de Proveedores de Servicios de Terceros", 0.6, 0.7, 0.42],
    ["Incompatibilidad con Nuevas Actualizaciones del Sistema Operativo", 0.5, 0.9, 0.45],
    ["Falta de Conexión a Internet durante las Pruebas de Geolocalización en Campo", 0.5, 0.6, 0.30],
    ["Saturación de la Base de Datos por Aumento de Usuarios Concurrentes", 0.5, 0.7, 0.35],
    ["Incumplimiento de Normativas de Privacidad en la Recolección de Datos", 0.3, 0.8, 0.24],
    ["Interrupción del Servicio de la API de Google Maps", 0.3, 0.8, 0.24],
    ["Fallo en la Sincronización de Datos entre Dispositivos y Servidor", 0.5, 0.7, 0.35],
    ["Falta de Personal para Pruebas de Campo en Áreas Remotas", 0.6, 0.7, 0.42],
    ["Limitaciones en el Rendimiento debido a Restricciones de Hardware", 0.3, 0.8, 0.24],
    ["Problemas de Compatibilidad con Dispositivos de Diferentes Marcas y Modelos", 0.4, 0.6, 0.24],
    ["Retrasos en la Obtención de Permisos para el Uso de Datos Geográficos de Organismos Locales", 0.3, 0.7, 0.21],
    ["Errores en la Traducción de Contenidos para Usuarios Internacionales", 0.4, 0.6, 0.24],
    ["Solicitudes de Cambios No Planificados por Parte de Stakeholders", 0.4, 0.7, 0.28],
    ["Retrasos en la Configuración de Servicios de Almacenamiento Local", 0.4, 0.6, 0.24]
]

# Sort data by exposure
risk_data_sorted = sorted(risk_data, key=lambda x: x[3], reverse=True)

# Create a Word document
doc = Document()
doc.add_heading('Ranking de Riesgos por Exposición', level=1)

# Create table with header
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Ranking'
hdr_cells[1].text = 'Riesgo'
hdr_cells[2].text = 'Probabilidad'
hdr_cells[3].text = 'Impacto'
hdr_cells[4].text = 'Exposición'

# Populate table with sorted data
for idx, (risk, probability, impact, exposure) in enumerate(risk_data_sorted, 1):
    row_cells = table.add_row().cells
    row_cells[0].text = str(idx)
    row_cells[1].text = risk
    row_cells[2].text = str(probability)
    row_cells[3].text = str(impact)
    row_cells[4].text = str(exposure)

# Save document
doc.save("Riesgos_Ranking_Exposicion_Sorted.docx")
