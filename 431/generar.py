from docx import Document
from docx.shared import Inches

# Crear el documento
doc = Document()

# Añadir un título
doc.add_heading('Comparación de MSF con PMBOK, CMMI y ISO 21500', level=1)

# Añadir una pequeña introducción
doc.add_paragraph(
    'En este documento se compara el Microsoft Solutions Framework (MSF) con tres estándares: PMBOK, CMMI, y ISO 21500. '
    'La tabla siguiente muestra una comparación de diversos aspectos como el país de origen, la organización que los respalda, '
    'certificaciones, aplicabilidad en Perú y áreas de conocimiento abordadas.'
)

# Crear la tabla
table = doc.add_table(rows=1, cols=5)
table.autofit = True

# Definir los encabezados de la tabla
headers = ['Aspecto', 'MSF', 'PMBOK', 'CMMI', 'ISO 21500']
hdr_cells = table.rows[0].cells
for i in range(5):
    hdr_cells[i].text = headers[i]

# Datos para cada fila de la tabla
data = [
    ['Nombre Completo', 'Microsoft Solutions Framework', 'Project Management Body of Knowledge', 
     'Capability Maturity Model Integration', 'ISO 21500:2012 – Guía sobre la gestión de proyectos'],
    ['País de Origen', 'Estados Unidos', 'Estados Unidos', 'Estados Unidos', 'Internacional (ISO, Suiza)'],
    ['Organización que lo Respaldó', 'Microsoft', 'Project Management Institute (PMI)', 
     'CMMI Institute (anteriormente SEI)', 'International Organization for Standardization (ISO)'],
    ['Certificaciones (Personas/Empresas)', 'No hay certificaciones formales', 
     'PMP, CAPM para personas, PgMP para empresas', 
     'Certificaciones para organizaciones en niveles de madurez (CMMI Dev, Svc, etc.)', 
     'No tiene certificaciones específicas, pero influye en otras certificaciones ISO'],
    ['Ejemplo de Certificación en Perú', 'No aplica', 'PMP y CAPM son populares en Perú', 
     'CMMI Dev es utilizado por empresas de software y tecnología en Perú', 
     'Implementado en procesos ISO en Perú'],
    ['Áreas de Conocimiento', 'Gestión de riesgos, calidad, desarrollo ágil, metodologías iterativas', 
     'Integración, alcance, tiempo, costo, calidad, recursos humanos, comunicación, riesgo, adquisiciones, interesados', 
     'Desempeño de procesos, gestión de proyectos en niveles de madurez', 
     'Integración, cronograma, costos, calidad, recursos, riesgos, adquisiciones'],
    ['Enfoque', 'Proyectos tecnológicos y de software', 'Gestión de proyectos en general', 
     'Mejorar la eficiencia y madurez de los procesos organizacionales', 
     'Gestión general de proyectos en cualquier industria'],
    ['Nivel de Formalización', 'Menos estructurado, adaptable', 'Altamente estructurado', 
     'Altamente formalizado, niveles de madurez', 'Estructurado pero flexible'],
    ['Evolución', 'Iniciado en 1994, evolucionado para metodologías ágiles y DevOps', 
     'Actualizado regularmente desde 1996, actualmente en su 7ª edición', 
     'Continuamente actualizado, enfocado en desarrollo y servicios', 
     'Introducido en 2012, alineado con otros estándares ISO'],
    ['Escalabilidad', 'Ideal para proyectos pequeños o medianos', 'Adecuado para proyectos de cualquier tamaño', 
     'Enfocado en grandes organizaciones', 'Aplicable a proyectos de cualquier tamaño o industria'],
    ['Certificaciones Dirigidas a', 'Principalmente empresas de software', 'Profesionales y empresas', 
     'Empresas que buscan mejorar la eficiencia interna', 'Compatible con certificaciones ISO'],
    ['Popularidad en Perú', 'Limitada', 'Amplio uso en consultorías y construcción', 
     'Adoptado en empresas tecnológicas y de software', 'Algunas organizaciones lo implementan como base ISO']
]

# Añadir los datos a la tabla
for row_data in data:
    row_cells = table.add_row().cells
    for i, cell_data in enumerate(row_data):
        row_cells[i].text = cell_data

# Guardar el documento
doc.save('comparacion_MSF_PMBO_CMMI_ISO21500.docx')

print("Documento generado exitosamente.")
