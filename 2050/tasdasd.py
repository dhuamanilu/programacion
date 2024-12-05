from docx import Document

# Crear un nuevo documento
doc = Document()
doc.add_heading('Comparativa entre Amazon Appstore, Samsung Galaxy Store y Google Play Store', 0)

# Agregar la tabla
table = doc.add_table(rows=1, cols=4)

# Establecer los encabezados de la tabla
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Característica'
hdr_cells[1].text = 'Amazon Appstore'
hdr_cells[2].text = 'Samsung Galaxy Store'
hdr_cells[3].text = 'Google Play Store'

# Definir los datos para las filas
rows = [
    ('Alcance Global', 'Presencia limitada, principalmente en EE.UU. y Europa', 'Presencia global, especialmente en dispositivos Samsung', 'Global, con presencia en más de 190 países'),
    ('Número de Aplicaciones', 'Aproximadamente 500,000', 'Más de 300,000', 'Más de 3 millones'),
    ('Política de Publicación', 'Revisión manual, pero menos estricta que Google', 'Revisión manual, similar a otras tiendas grandes', 'Estricta, con revisión manual y automática'),
    ('Soporte para Dispositivos', 'Android, Fire OS (dispositivos Amazon)', 'Android, dispositivos Samsung (smartphones, tablets, wearables)', 'Android, Wear OS, Android TV, Chrome OS'),
    ('Modelo de Monetización', 'Pago por descarga, compras dentro de la app, suscripciones', 'Pago por descarga, compras dentro de la app, suscripciones', 'Pago por descarga, compras dentro de la app, suscripciones'),
    ('Comisiones de la Tienda', '20% (típicamente más baja)', '30% (comisión estándar en la mayoría de aplicaciones)', '15% - 30% (dependiendo del tipo de app y las ganancias)'),
    ('Interfaz de Usuario', 'Interfaz más simple, optimizada para dispositivos Amazon', 'Interfaz fluida, integrada con los servicios Samsung', 'Interfaz bien establecida y optimizada'),
    ('Requerimientos para Desarrolladores', 'Requiere cuenta de desarrollador gratuita o con tarifa de entrada', 'Requiere registro como desarrollador y aceptación de políticas de Samsung', 'Requiere cuenta de desarrollador con tarifa anual (25 USD)'),
    ('Actualizaciones de Aplicaciones', 'Automáticas, con énfasis en dispositivos Amazon', 'Automáticas a través de Samsung Galaxy Store y actualizaciones OTA', 'Automáticas a través de Google Play Services'),
    ('Seguridad', 'Seguridad aceptable, revisión de aplicaciones enfocada en evitar malware', 'Alta seguridad, integrada con Knox Security de Samsung', 'Alto nivel de seguridad, pero ocasionalmente se encuentran apps maliciosas'),
    ('Soporte para In-App Purchases', 'Totalmente integrado con Amazon In-App Purchasing', 'Totalmente integrado con Samsung In-App Purchasing', 'Totalmente integrado con Google Play Billing'),
    ('Sistemas Operativos Soportados', 'Android, Fire OS', 'Android (principalmente en dispositivos Samsung)', 'Android, Wear OS, Chrome OS')
]

# Agregar las filas a la tabla
for row in rows:
    row_cells = table.add_row().cells
    for i, cell_value in enumerate(row):
        row_cells[i].text = cell_value

# Guardar el documento
doc.save('comparativa_appstores_google.docx')

print("El documento 'comparativa_appstores_google.docx' ha sido generado exitosamente.")
